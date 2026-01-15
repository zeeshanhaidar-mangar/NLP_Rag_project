"""
DocuMind AI - Intelligent Document Q&A Assistant
Complete RAG system with 30+ features
"""

import streamlit as st
import os
from datetime import datetime
import re

# Core imports
try:
    from sentence_transformers import SentenceTransformer
    import faiss
    import numpy as np
    from PyPDF2 import PdfReader
    import google.generativeai as genai
    from docx import Document as DocxDocument
    import pandas as pd
    import pdfplumber
except ImportError as e:
    st.error(f"⏳ Installing dependencies... Refresh in 1 minute.")
    st.stop()

# Page config
st.set_page_config(
    page_title="DocuMind AI",
    page_icon="📚",
    layout="wide"
)

# CSS
st.markdown("""
<style>
.main-header {
    font-size: 3rem;
    font-weight: bold;
    background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    text-align: center;
    padding: 1rem;
}
.stButton>button {
    width: 100%;
    border-radius: 8px;
    height: 3em;
    font-weight: 600;
}
</style>
""", unsafe_allow_html=True)

# Session state
def init_session_state():
    defaults = {
        'initialized': False,
        'embedding_model': None,
        'document_store': [],
        'metadata_store': [],
        'faiss_index': None,
        'question_history': [],
        'api_configured': False,
        'api_key': '',
        'stats': {
            'total_documents': 0,
            'total_chunks': 0,
            'total_queries': 0,
            'avg_confidence': 0.0
        }
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# Load model
@st.cache_resource(show_spinner=False)
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

# Clean text
def clean_text(text):
    if not text:
        return ""
    text = text.replace('\x00', '')
    text = ''.join(c for c in text if c.isprintable() or c in ['\n', '\t'])
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

# Chunk text
def chunk_text(text, size=500, overlap=50):
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunk = text[start:end].strip()
        if len(chunk) >= 50:
            chunks.append(chunk)
        start += size - overlap
    return chunks

# Process PDF
def process_pdf(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        if text.strip():
            return text, "success", len(reader.pages)
    except:
        pass
    try:
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
            return text, "success", len(pdf.pages)
    except Exception as e:
        return "", f"error: {e}", 0

# Process DOCX
def process_docx(file):
    try:
        doc = DocxDocument(file)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return text, "success", len(doc.paragraphs)
    except Exception as e:
        return "", f"error: {e}", 0

# Process TXT
def process_txt(file):
    for enc in ['utf-8', 'latin-1', 'cp1252']:
        try:
            file.seek(0)
            text = file.read().decode(enc)
            return text, "success", len(text)
        except:
            continue
    return "", "error: encoding", 0

# Process CSV
def process_csv(file):
    try:
        df = pd.read_csv(file)
        text = f"CSV: {len(df)} rows\nColumns: {', '.join(df.columns)}\n\n"
        text += df.to_string()
        return text, "success", len(df)
    except Exception as e:
        return "", f"error: {e}", 0

# Main processor
def process_document(file):
    ext = file.name.split('.')[-1].lower()
    processors = {
        'pdf': process_pdf,
        'docx': process_docx,
        'txt': process_txt,
        'csv': process_csv
    }
    processor = processors.get(ext)
    if processor:
        return processor(file)
    return "", f"Unsupported: {ext}", 0

# Generate summary
def generate_summary(text, filename):
    try:
        if not st.session_state.api_configured:
            return "API not configured"
        prompt = f"Summarize in 2-3 sentences:\n{text[:2000]}"
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Summary unavailable: {str(e)}"

# Search
def search(query, k=5):
    if not st.session_state.faiss_index or not st.session_state.document_store:
        return [], [], []
    
    q_emb = st.session_state.embedding_model.encode([query]).astype('float32')
    distances, indices = st.session_state.faiss_index.search(q_emb, k)
    
    max_dist = np.max(distances[0]) if len(distances[0]) > 0 else 1
    similarities = 1 - (distances[0] / (max_dist + 1e-6))
    
    chunks = [st.session_state.document_store[i] for i in indices[0] if i < len(st.session_state.document_store)]
    metadata = [st.session_state.metadata_store[i] for i in indices[0] if i < len(st.session_state.metadata_store)]
    scores = [float(s) for s in similarities]
    
    return chunks, metadata, scores

# Generate answer
def generate_answer(query, chunks, metadata):
    try:
        if not st.session_state.api_configured:
            return "⚠️ Configure API key in sidebar", 0
        
        if not chunks:
            return "❌ No relevant info found. Upload documents first.", 0
        
        context = "\n\n".join([f"[{m['source']}]\n{c}" for c, m in zip(chunks, metadata)])
        
        prompt = f"""Answer based ONLY on context. Cite sources.

Context:
{context}

Question: {query}

Answer:"""
        
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(prompt)
        
        confidence = np.mean([m.get('score', 0) for m in metadata]) * 100
        
        return response.text, confidence
    except Exception as e:
        return f"Error generating answer: {str(e)}", 0

# Sidebar
with st.sidebar:
    st.markdown("# 📚 DocuMind AI")
    st.markdown("---")
    
    st.markdown("### 🔑 API Configuration")
    with st.expander("ℹ️ How to get API key"):
        st.markdown("""
        1. Visit [Google AI Studio](https://makersuite.google.com/app/apikey)
        2. Click "Create API Key"
        3. Enable "Generative Language API" if prompted
        4. Copy the key and paste below
        """)
    
    api_key = st.text_input(
        "Gemini API Key", 
        type="password", 
        value=st.session_state.api_key,
        placeholder="AIza..."
    )
    
    if api_key and api_key != st.session_state.api_key:
        if st.button("🔌 Configure API", type="primary"):
            with st.spinner("Validating API key..."):
                try:
                    # Configure the API
                    genai.configure(api_key=api_key)
                    
                    # Test with a simple request
                    model = genai.GenerativeModel('gemini-pro')
                    response = model.generate_content("Say 'API key is working' if you receive this message")
                    
                    # Check if we got a valid response
                    if response and response.text:
                        st.session_state.api_key = api_key
                        st.session_state.api_configured = True
                        st.success("✅ API key validated successfully!")
                        st.balloons()
                        st.rerun()
                    else:
                        st.error("❌ Invalid response from API")
                        
                except Exception as e:
                    error_message = str(e)
                    st.error(f"❌ API key validation failed")
                    
                    # Provide helpful error messages
                    if "API_KEY_INVALID" in error_message or "invalid" in error_message.lower():
                        st.warning("**Possible issues:**")
                        st.write("• API key is incorrect or incomplete")
                        st.write("• Copy the full key starting with 'AIza'")
                    elif "PERMISSION_DENIED" in error_message:
                        st.warning("**Possible issues:**")
                        st.write("• Enable 'Generative Language API' in Google Cloud Console")
                        st.write("• Check if billing is enabled on your project")
                    else:
                        st.warning(f"**Error details:** {error_message}")
                    
                    st.info("💡 **Troubleshooting:**\n- Ensure API key starts with 'AIza'\n- Enable Generative Language API\n- Check Google Cloud billing")
    
    if st.session_state.api_configured:
        st.success("✅ API Connected")
        if st.button("🔄 Reset API"):
            st.session_state.api_key = ''
            st.session_state.api_configured = False
            st.rerun()
    
    st.markdown("---")
    page = st.radio("📍 Navigation", ["🏠 Home", "📤 Upload", "❓ Ask", "📊 Stats"])
    
    st.markdown("---")
    st.markdown("### 📈 Quick Stats")
    st.metric("Documents", st.session_state.stats['total_documents'])
    st.metric("Chunks", st.session_state.stats['total_chunks'])
    st.metric("Queries", st.session_state.stats['total_queries'])

# Initialize model
if not st.session_state.initialized:
    with st.spinner("🔄 Loading AI models..."):
        st.session_state.embedding_model = load_model()
        st.session_state.initialized = True

# HOME PAGE
if page == "🏠 Home":
    st.markdown('<div class="main-header">📚 DocuMind AI</div>', unsafe_allow_html=True)
    st.markdown("### *Intelligent Document Q&A Assistant*")
    
    st.markdown("---")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("#### 📤 Upload Documents")
        st.write("• PDF, DOCX, TXT, CSV")
        st.write("• Automatic text extraction")
        st.write("• Smart chunking")
        st.write("• Multi-format support")
    
    with col2:
        st.markdown("#### 🤖 AI-Powered Answers")
        st.write("• Natural language queries")
        st.write("• Source citations")
        st.write("• Confidence scores")
        st.write("• Context-aware responses")
    
    with col3:
        st.markdown("#### 📊 Advanced Analytics")
        st.write("• Query history tracking")
        st.write("• Usage statistics")
        st.write("• Performance metrics")
        st.write("• Document insights")
    
    st.markdown("---")
    st.markdown("### 🚀 Quick Start Guide")
    
    with st.expander("**Step 1: Configure API Key**", expanded=not st.session_state.api_configured):
        st.markdown("""
        1. Visit [Google AI Studio](https://makersuite.google.com/app/apikey)
        2. Sign in with your Google account
        3. Click "Create API Key"
        4. Copy the generated key
        5. Paste it in the sidebar and click "Configure API"
        """)
    
    with st.expander("**Step 2: Upload Documents**"):
        st.markdown("""
        1. Navigate to the "📤 Upload" page
        2. Click "Choose files" and select your documents
        3. Supported formats: PDF, DOCX, TXT, CSV
        4. Click "🚀 Process" to analyze documents
        5. Wait for processing to complete
        """)
    
    with st.expander("**Step 3: Ask Questions**"):
        st.markdown("""
        1. Go to the "❓ Ask" page
        2. Type your question in natural language
        3. Click "🔍 Get Answer"
        4. Review the AI-generated answer with sources
        5. Check confidence score for answer reliability
        """)
    
    st.markdown("---")
    st.markdown("### ℹ️ Features")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        **Document Processing:**
        - Multi-format support
        - Automatic text cleaning
        - Intelligent chunking
        - Metadata extraction
        
        **Search & Retrieval:**
        - Semantic search
        - Vector embeddings
        - FAISS indexing
        - Top-k retrieval
        """)
    
    with col2:
        st.markdown("""
        **AI Capabilities:**
        - Gemini AI integration
        - Context-aware answers
        - Source attribution
        - Confidence scoring
        
        **Analytics:**
        - Question history
        - Usage statistics
        - Performance tracking
        - Document analytics
        """)

# UPLOAD PAGE
elif page == "📤 Upload":
    st.title("📤 Upload Documents")
    
    st.markdown("""
    Upload your documents to build a searchable knowledge base. 
    Supported formats: **PDF, DOCX, TXT, CSV**
    """)
    
    st.markdown("---")
    
    files = st.file_uploader(
        "Choose files to upload", 
        type=['pdf', 'docx', 'txt', 'csv'], 
        accept_multiple_files=True,
        help="Select one or more documents to process"
    )
    
    if files:
        st.info(f"📁 {len(files)} file(s) selected")
    
    if st.button("🚀 Process Documents", type="primary", disabled=not files):
        if not files:
            st.error("❌ Please select files to upload")
        else:
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            total_processed = 0
            total_failed = 0
            
            for i, file in enumerate(files):
                status_text.text(f"Processing {i+1}/{len(files)}: {file.name}")
                
                with st.expander(f"📄 {file.name}", expanded=True):
                    file_status = st.empty()
                    file_status.info("⏳ Processing...")
                    
                    try:
                        # Process document
                        text, result, info = process_document(file)
                        
                        if "error" in result:
                            file_status.error(f"❌ Failed: {result}")
                            total_failed += 1
                            continue
                        
                        # Clean text
                        text = clean_text(text)
                        
                        if len(text) < 100:
                            file_status.warning("⚠️ Document too short (< 100 characters)")
                            total_failed += 1
                            continue
                        
                        # Generate summary
                        summary = generate_summary(text, file.name)
                        
                        # Create chunks
                        chunks = chunk_text(text)
                        
                        if not chunks:
                            file_status.warning("⚠️ No valid chunks created")
                            total_failed += 1
                            continue
                        
                        # Generate embeddings
                        embeddings = st.session_state.embedding_model.encode(chunks).astype('float32')
                        
                        # Initialize or add to FAISS index
                        if st.session_state.faiss_index is None:
                            st.session_state.faiss_index = faiss.IndexFlatL2(384)
                        
                        st.session_state.faiss_index.add(embeddings)
                        st.session_state.document_store.extend(chunks)
                        
                        # Add metadata
                        metadata = [{'source': file.name, 'score': 0} for _ in chunks]
                        st.session_state.metadata_store.extend(metadata)
                        
                        # Update statistics
                        st.session_state.stats['total_documents'] += 1
                        st.session_state.stats['total_chunks'] += len(chunks)
                        
                        # Display success
                        file_status.success("✅ Successfully processed!")
                        st.write(f"**Chunks created:** {len(chunks)}")
                        st.write(f"**Text length:** {len(text):,} characters")
                        st.write(f"**Summary:** {summary}")
                        
                        total_processed += 1
                    
                    except Exception as e:
                        file_status.error(f"❌ Error: {str(e)}")
                        total_failed += 1
                
                # Update progress
                progress_bar.progress((i + 1) / len(files))
            
            # Final status
            status_text.text("Processing complete!")
            
            if total_processed > 0:
                st.balloons()
                st.success(f"""
                ✅ **Processing Complete!**
                - Successfully processed: {total_processed} documents
                - Failed: {total_failed} documents
                - Total documents in system: {st.session_state.stats['total_documents']}
                - Total chunks: {st.session_state.stats['total_chunks']}
                """)
            else:
                st.error("❌ No documents were processed successfully")

# ASK PAGE
elif page == "❓ Ask":
    st.title("❓ Ask Questions")
    
    # Check API configuration
    if not st.session_state.api_configured:
        st.error("⚠️ **API key not configured**")
        st.info("Please configure your Gemini API key in the sidebar to use this feature.")
        st.stop()
    
    # Check if documents are uploaded
    if not st.session_state.document_store:
        st.warning("⚠️ **No documents uploaded**")
        st.info("Please upload documents first using the Upload page.")
        st.stop()
    
    st.markdown(f"""
    Ask questions about your uploaded documents. The AI will search through 
    **{st.session_state.stats['total_documents']} documents** and 
    **{st.session_state.stats['total_chunks']} chunks** to find relevant information.
    """)
    
    st.markdown("---")
    
    # Question input
    question = st.text_area(
        "Your Question:", 
        placeholder="What would you like to know about your documents?",
        height=100,
        help="Type your question in natural language"
    )
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        ask_button = st.button("🔍 Get Answer", type="primary", use_container_width=True)
    
    with col2:
        num_sources = st.selectbox("Sources", [3, 5, 7, 10], index=1)
    
    if ask_button:
        if len(question.strip()) < 3:
            st.error("❌ Please enter a question (at least 3 characters)")
        else:
            with st.spinner("🤔 Analyzing documents and generating answer..."):
                start_time = datetime.now()
                
                # Search for relevant chunks
                chunks, metadata, scores = search(question, k=num_sources)
                
                if not chunks:
                    st.warning("⚠️ No relevant information found in uploaded documents")
                    st.stop()
                
                # Add scores to metadata
                for i, score in enumerate(scores):
                    if i < len(metadata):
                        metadata[i]['score'] = score
                
                # Generate answer
                answer, confidence = generate_answer(question, chunks, metadata)
                
                # Calculate time taken
                time_taken = (datetime.now() - start_time).total_seconds()
                
                # Save to history
                st.session_state.question_history.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'question': question,
                    'answer': answer,
                    'confidence': confidence,
                    'time_taken': time_taken
                })
                
                st.session_state.stats['total_queries'] += 1
                
                # Display results
                st.markdown("---")
                st.markdown("### 💬 Answer")
                
                # Confidence indicator
                if confidence >= 70:
                    st.success(f"🟢 High Confidence: {confidence:.1f}%")
                elif confidence >= 40:
                    st.warning(f"🟡 Medium Confidence: {confidence:.1f}%")
                else:
                    st.error(f"🔴 Low Confidence: {confidence:.1f}%")
                
                # Answer
                st.markdown(answer)
                
                # Metadata
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("⏱️ Response Time", f"{time_taken:.2f}s")
                with col2:
                    st.metric("📚 Sources Used", len(set([m['source'] for m in metadata])))
                
                # Sources section
                st.markdown("---")
                st.markdown("### 📚 Sources")
                
                sources_with_scores = {}
                for i, m in enumerate(metadata):
                    source = m['source']
                    score = m.get('score', 0)
                    if source not in sources_with_scores:
                        sources_with_scores[source] = []
                    sources_with_scores[source].append((score, chunks[i]))
                
                for idx, (source, score_chunk_pairs) in enumerate(sorted(sources_with_scores.items()), 1):
                    avg_score = np.mean([sc[0] for sc in score_chunk_pairs]) * 100
                    with st.expander(f"{idx}. **{source}** (Relevance: {avg_score:.1f}%)"):
                        st.write(f"**Relevant excerpts found:** {len(score_chunk_pairs)}")
                        for i, (score, chunk) in enumerate(score_chunk_pairs[:3], 1):
                            st.markdown(f"**Excerpt {i}:** {chunk[:200]}...")

# STATS PAGE
elif page == "📊 Stats":
    st.title("📊 Statistics & Analytics")
    
    st.markdown("---")
    
    # Overall statistics
    st.markdown("### 📈 Overall Statistics")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric(
            "📄 Documents", 
            st.session_state.stats['total_documents'],
            help="Total documents uploaded"
        )
    
    with col2:
        st.metric(
            "🧩 Chunks", 
            st.session_state.stats['total_chunks'],
            help="Total text chunks indexed"
        )
    
    with col3:
        st.metric(
            "❓ Queries", 
            st.session_state.stats['total_queries'],
            help="Total questions asked"
        )
    
    with col4:
        avg_conf = np.mean([q['confidence'] for q in st.session_state.question_history]) if st.session_state.question_history else 0
        st.metric(
            "📊 Avg Confidence", 
            f"{avg_conf:.1f}%",
            help="Average confidence score"
        )
    
    st.markdown("---")
    
    # Document breakdown
    if st.session_state.metadata_store:
        st.markdown("### 📚 Document Breakdown")
        
        doc_counts = {}
        for meta in st.session_state.metadata_store:
            source = meta['source']
            doc_counts[source] = doc_counts.get(source, 0) + 1
        
        df_docs = pd.DataFrame([
            {'Document': k, 'Chunks': v} 
            for k, v in sorted(doc_counts.items(), key=lambda x: x[1], reverse=True)
        ])
        
        st.dataframe(df_docs, use_container_width=True)
    
    st.markdown("---")
    
    # Question history
    st.markdown("### 📜 Question History")
    
    if st.session_state.question_history:
        # Show recent questions
        num_to_show = st.slider("Number of recent questions to display", 5, 20, 10)
        
        for i, entry in enumerate(reversed(st.session_state.question_history[-num_to_show:]), 1):
            with st.expander(f"**Q{i}:** {entry['question'][:80]}{'...' if len(entry['question']) > 80 else ''}"):
                st.write(f"**⏰ Time:** {entry['timestamp']}")
                st.write(f"**📊 Confidence:** {entry['confidence']:.1f}%")
                st.write(f"**⏱️ Response Time:** {entry.get('time_taken', 'N/A'):.2f}s" if 'time_taken' in entry else "")
                st.write(f"**💬 Answer:**")
                st.write(entry['answer'][:500] + ('...' if len(entry['answer']) > 500 else ''))
        
        # Export history
        if st.button("📥 Export History as CSV"):
            df_history = pd.DataFrame(st.session_state.question_history)
            csv = df_history.to_csv(index=False)
            st.download_button(
                "Download CSV",
                csv,
                "question_history.csv",
                "text/csv",
                key='download-csv'
            )
    else:
        st.info("💡 No questions asked yet. Go to the Ask page to start querying your documents!")
    
    st.markdown("---")
    
    # Reset options
    st.markdown("### ⚙️ Settings")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🗑️ Clear Question History", type="secondary"):
            st.session_state.question_history = []
            st.session_state.stats['total_queries'] = 0
            st.success("✅ Question history cleared!")
            st.rerun()
    
    with col2:
        if st.button("🔄 Reset All Data", type="secondary"):
            if st.session_state.stats['total_documents'] > 0:
                confirm = st.checkbox("⚠️ Confirm reset (this will delete all documents)")
                if confirm:
                    for key in ['document_store', 'metadata_store', 'faiss_index', 'question_history']:
                        if key == 'faiss_index':
                            st.session_state[key] = None
                        else:
                            st.session_state[key] = []
                    st.session_state.stats = {
                        'total_documents': 0,
                        'total_chunks': 0,
                        'total_queries': 0,
                        'avg_confidence': 0.0
                    }
                    st.success("✅ All data reset!")
                    st.rerun()
            else:
                st.info("No data to reset")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray; padding: 1rem;'>
    <p>DocuMind AI © 2024 | Powered by Gemini AI & Streamlit</p>
    <p style='font-size: 0.8rem;'>Intelligent Document Q&A Assistant with RAG Technology</p>
</div>
""", unsafe_allow_html=True)
